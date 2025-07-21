"""
Word文档解析器

基于python-docx库解析Word文档(.docx格式)。
"""

import os
from typing import Dict, Any, Optional
from .base_parser import BaseParser

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordParser(BaseParser):
    """Word文档解析器
    
    支持解析.docx格式的Word文档，提取文本内容。
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """初始化Word解析器
        
        Args:
            config: 配置参数
        """
        super().__init__(config)
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for WordParser. "
                "Install it with: pip install python-docx"
            )
    
    def _get_supported_extensions(self) -> list:
        """获取支持的文件扩展名"""
        return ['.docx']
    
    def _parse_file_content(self, file_path: str) -> str:
        """解析Word文档内容
        
        Args:
            file_path: Word文档路径
            
        Returns:
            解析后的文本内容
            
        Raises:
            FileNotFoundError: 文件不存在
            Exception: 解析失败
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        try:
            # 使用python-docx解析Word文档
            doc = Document(file_path)
            
            # 提取所有段落的文本
            content = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # 跳过空段落
                    content.append(text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        content.append(" | ".join(row_text))
            
            return "\n".join(content)
            
        except Exception as e:
            raise Exception(f"解析Word文档失败 {file_path}: {str(e)}")
    
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """获取Word文档元数据
        
        Args:
            file_path: Word文档路径
            
        Returns:
            文档元数据
        """
        metadata = super().get_metadata(file_path)
        
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            # 添加Word特有的元数据
            metadata.update({
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": core_props.created.isoformat() if core_props.created else "",
                "modified": core_props.modified.isoformat() if core_props.modified else "",
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables)
            })
            
        except Exception as e:
            metadata["error"] = f"获取元数据失败: {str(e)}"
        
        return metadata
