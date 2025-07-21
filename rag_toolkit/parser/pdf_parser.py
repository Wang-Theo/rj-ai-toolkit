"""
PDF文档解析器

支持PDF文件的文本提取和结构化解析。
"""

from typing import List, Optional
from pathlib import Path
from .base_parser import BaseParser, DocumentMetadata
import fitz  # PyMuPDF


class PDFParser(BaseParser):
    """PDF文档解析器"""
    
    def _get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名"""
        return ["pdf"]
    
    def _parse_file_content(self, file_path: Path) -> str:
        """解析PDF文件内容"""
        try:
            doc = fitz.open(str(file_path))
            text_content = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                
                if self.config.preserve_structure:
                    text_content.append(f"\n--- 第 {page_num + 1} 页 ---\n{text}")
                else:
                    text_content.append(text)
            
            doc.close()
            full_text = "\n".join(text_content)
            
            return self._clean_text(full_text)
            
        except Exception as e:
            raise ValueError(f"解析PDF文件失败: {e}")
    
    def _extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """提取PDF元数据"""
        metadata = super()._extract_metadata(file_path)
        
        try:
            doc = fitz.open(str(file_path))
            pdf_metadata = doc.metadata
            
            # 添加PDF特有的元数据
            metadata.page_count = len(doc)
            metadata.author = pdf_metadata.get('author', '')
            metadata.title = pdf_metadata.get('title', '')
            
            doc.close()
            
        except Exception:
            # 如果无法提取PDF元数据，使用默认值
            pass
        
        return metadata


class WordParser(BaseParser):
    """Word文档解析器"""
    
    def _get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名"""
        return ["docx", "doc"]
    
    def _parse_file_content(self, file_path: Path) -> str:
        """解析Word文件内容"""
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # 提取表格内容
            if self.config.extract_tables:
                for table in doc.tables:
                    table_text = self._extract_table_text(table)
                    if table_text:
                        text_content.append(f"\n[表格]\n{table_text}\n")
            
            full_text = "\n".join(text_content)
            return self._clean_text(full_text)
            
        except ImportError:
            raise ValueError("需要安装 python-docx 库来解析Word文档: pip install python-docx")
        except Exception as e:
            raise ValueError(f"解析Word文件失败: {e}")
    
    def _extract_table_text(self, table) -> str:
        """提取表格文本"""
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(" | ".join(row_data))
        
        return "\n".join(table_data)
    
    def _extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """提取Word元数据"""
        metadata = super()._extract_metadata(file_path)
        
        try:
            from docx import Document
            doc = Document(str(file_path))
            
            # 提取Word文档属性
            core_properties = doc.core_properties
            metadata.author = core_properties.author or ""
            metadata.title = core_properties.title or ""
            
            # 估算页数（每页约500字）
            total_chars = sum(len(p.text) for p in doc.paragraphs)
            metadata.page_count = max(1, total_chars // 500)
            
        except Exception:
            pass
        
        return metadata
